# -*- coding: utf-8 -*-
"""操作マニュアル作成スクリプト

中学生でも読みやすいカジュアルな操作マニュアルをWord形式で作成する
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image, ImageDraw

# 出力先
OUTPUT_DIR = Path(__file__).parent.parent / "docs"
IMAGES_DIR = OUTPUT_DIR / "manual_images"
OUTPUT_FILE = OUTPUT_DIR / "操作マニュアル_動画背景除去ツール.docx"


def create_placeholder_images():
    """プレースホルダー画像を作成"""
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    width, height = 500, 350

    screens = [
        ("01_initial.png", "初期画面", "#FAFAFA", [
            "📁",
            "",
            "動画をドラッグ＆ドロップ",
            "または クリック",
            "",
            ".mp4  .mov  .m4v"
        ]),
        ("02_selected.png", "動画を選んだあと", "#FAFAFA", [
            "[サムネイル]",
            "",
            "sample_video.mp4",
            "7分23秒",
            "",
            "[🚀 背景を除去する]"
        ]),
        ("03_processing.png", "処理中", "#FAFAFA", [
            "[サムネイル]",
            "",
            "67%",
            "8,902 / 13,290",
            "",
            "[キャンセル]"
        ]),
        ("04_done.png", "完了！", "#FAFAFA", [
            "[背景なしサムネイル]",
            "",
            "✅ 完了!",
            "",
            "[💾 ファイルを保存]"
        ]),
    ]

    for filename, title, bg_color, lines in screens:
        img = Image.new("RGB", (width, height), bg_color)
        draw = ImageDraw.Draw(img)

        # タイトルバー
        draw.rectangle([0, 0, width, 35], fill="#8BC34A")

        # 枠線
        draw.rectangle([0, 0, width-1, height-1], outline="#E0E0E0", width=2)

        # コンテンツ
        y = 60
        for line in lines:
            x = width // 2 - len(line) * 7
            draw.text((max(20, x), y), line, fill="#263238")
            y += 35

        # ラベル
        draw.rectangle([10, height-40, width-10, height-10], fill="#263238")
        draw.text((20, height-32), title, fill="#FFFFFF")

        img.save(IMAGES_DIR / filename)
        print(f"Created: {filename}")


def create_manual():
    """カジュアルで読みやすいマニュアルを作成"""
    doc = Document()

    # スタイル設定
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(12)

    # ===== 表紙 =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    run = title.add_run("🎬 動画背景除去ツール")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0xC3, 0x4A)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("かんたん使い方ガイド")
    run.font.size = Pt(20)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    org = doc.add_paragraph()
    run = org.add_run("META AI LABO")
    run.font.size = Pt(14)
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== このアプリでできること =====
    h = doc.add_heading("🌟 このアプリでできること", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x8B, 0xC3, 0x4A)

    intro = """
動画の背景を「消す」ことができるアプリです！

たとえば...
"""
    doc.add_paragraph(intro.strip())

    examples = [
        "✨ 自分が映っている動画から、背景だけを消せる",
        "✨ グリーンバック（緑の布）がなくても大丈夫",
        "✨ 消した動画は、別の背景と合成できる",
    ]
    for ex in examples:
        p = doc.add_paragraph(ex)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph()

    use_case = doc.add_paragraph()
    use_case.add_run("📹 こんなときに便利！").bold = True
    doc.add_paragraph("・YouTubeやTikTokの動画編集")
    doc.add_paragraph("・プレゼン動画の作成")
    doc.add_paragraph("・ゲーム実況の顔出し合成")

    doc.add_page_break()

    # ===== 使い方（3ステップ）=====
    h = doc.add_heading("📱 使い方はカンタン3ステップ！", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x8B, 0xC3, 0x4A)

    doc.add_paragraph()

    # ステップ1
    step1 = doc.add_paragraph()
    run = step1.add_run("Step 1️⃣  動画ファイルを選ぶ")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x26, 0x32, 0x38)

    doc.add_paragraph()
    doc.add_paragraph("アプリを開いたら、この画面が出てきます。")

    img_path = IMAGES_DIR / "01_initial.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("動画ファイルを画面にドラッグ＆ドロップするか、")
    doc.add_paragraph("画面をクリックしてファイルを選んでください。")

    doc.add_paragraph()

    tip = doc.add_paragraph()
    tip.add_run("💡 ").font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
    tip.add_run("使える動画の種類: ").bold = True
    tip.add_run(".mp4、.mov、.m4v")

    doc.add_page_break()

    # ステップ2
    step2 = doc.add_paragraph()
    run = step2.add_run("Step 2️⃣  背景を除去する")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x26, 0x32, 0x38)

    doc.add_paragraph()
    doc.add_paragraph("動画を選ぶと、こんな画面になります。")

    img_path = IMAGES_DIR / "02_selected.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("緑色の「🚀 背景を除去する」ボタン").bold = True
    p.add_run("をクリック！")

    doc.add_paragraph()
    doc.add_paragraph("すると、AIが自動で背景を消してくれます。")

    img_path = IMAGES_DIR / "03_processing.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    wait = doc.add_paragraph()
    wait.add_run("⏳ ").font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
    wait.add_run("処理には少し時間がかかります。").bold = True
    doc.add_paragraph("動画の長さによって数分〜数十分かかることもあるので、")
    doc.add_paragraph("そのまま待っていてください。")

    doc.add_page_break()

    # ステップ3
    step3 = doc.add_paragraph()
    run = step3.add_run("Step 3️⃣  できた動画を保存する")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x26, 0x32, 0x38)

    doc.add_paragraph()
    doc.add_paragraph("処理が終わると「✅ 完了!」と表示されます。")

    img_path = IMAGES_DIR / "04_done.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("「💾 ファイルを保存」ボタン").bold = True
    p.add_run("をクリックして、")
    doc.add_paragraph("好きな場所に保存すれば完了です！")

    doc.add_paragraph()

    done = doc.add_paragraph()
    done.add_run("🎉 ").font.size = Pt(16)
    done.add_run("おめでとう！背景が透明な動画ができました！").bold = True

    doc.add_page_break()

    # ===== よくある質問 =====
    h = doc.add_heading("❓ よくある質問", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x8B, 0xC3, 0x4A)

    faqs = [
        ("処理にどのくらい時間がかかる？",
         "動画の長さとパソコンによって変わります。\n"
         "1分の動画で、だいたい1〜5分くらいです。\n"
         "長い動画だともっと時間がかかります。"),

        ("処理中、他のことしてもいい？",
         "大丈夫です！\n"
         "ただ、パソコンが少し重くなるかもしれません。\n"
         "電源は切らないでくださいね。"),

        ("できた動画はどうやって使うの？",
         "動画編集ソフト（iMovie、Premiere、DaVinciなど）で\n"
         "読み込むと、別の背景と合成できます！\n"
         "透明な部分に好きな画像や動画を入れられます。"),

        ("背景がうまく消えない...",
         "こんな動画だとキレイに消えやすいです：\n"
         "・明るい場所で撮った動画\n"
         "・背景がシンプルな動画\n"
         "・人がハッキリ映っている動画"),

        ("音声はどうなる？",
         "元の動画に音声があれば、そのまま残ります！\n"
         "音が消えることはないので安心してください。"),
    ]

    for q, a in faqs:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"Q. {q}").bold = True

        for line in a.split("\n"):
            doc.add_paragraph(f"　{line}")

    doc.add_page_break()

    # ===== 困ったときは =====
    h = doc.add_heading("🆘 困ったときは", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x8B, 0xC3, 0x4A)

    troubles = [
        ("アプリが開かない",
         "【Macの場合】\n"
         "初めて開くときは、アプリを右クリックして\n"
         "「開く」を選んでください。\n\n"
         "【Windowsの場合】\n"
         "セキュリティソフトがブロックしていることがあります。\n"
         "「詳細情報」→「実行」をクリックしてみてください。"),

        ("処理がすごく遅い",
         "パソコンのスペックによって速さが変わります。\n"
         "・他のアプリを閉じてみる\n"
         "・パソコンを再起動してみる\n"
         "これで少し早くなることがあります。"),

        ("エラーが出た",
         "・動画ファイルが壊れていないか確認\n"
         "・ディスクの空き容量を確認\n"
         "（動画サイズの3倍くらい必要です）"),
    ]

    for title, content in troubles:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"😥 {title}").bold = True

        for line in content.split("\n"):
            doc.add_paragraph(f"　{line}")

    # ===== おわりに =====
    doc.add_page_break()

    for _ in range(5):
        doc.add_paragraph()

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("🎬 動画背景除去ツール")
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("かんたん使い方ガイド")

    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("META AI LABO")
    footer.add_run("\n")
    footer.add_run("Ver. 1.0.0")

    # 保存
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\n✅ マニュアルを作成しました: {OUTPUT_FILE}")


def main():
    print("操作マニュアル作成スクリプト")
    print("=" * 50)

    # スクリーンショットが存在するか確認
    screenshots = [
        "01_initial.png", "02_selected.png",
        "03_processing.png", "04_done.png"
    ]
    all_exist = all((IMAGES_DIR / f).exists() for f in screenshots)

    if all_exist:
        print("\n実際のスクリーンショットを使用します")
    else:
        print("\n1. サンプル画像を作成中...")
        create_placeholder_images()

    print("\n2. Word文書を作成中...")
    create_manual()

    print("\n完了しました!")
    print(f"出力先: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
